"""Generate internship review from organization as .docx (1 page)"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_para(text, bold=False, align=None, space_after=0, indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    return p


# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('REVIEW')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
p.paragraph_format.space_after = Pt(2)

add_para(
    'of the work of the student of the Master\u2019s Programme '
    '"Business Analytics and Big Data Systems" of the Graduate School of Business, '
    'NRU HSE, during the Industrial Internship period',
    align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=6
)

add_para(
    'This review is given to Dementev Gleb Dmitrievich (Group BABDS-242), who completed '
    'the Industrial Internship at TargetAI LLC (Moscow) from 12 January to 13 March 2026 '
    '(9 weeks). The internship topic was "Optimizing Customer Service Automation with '
    'Human-in-the-Loop: A Framework for Supervised AI Agent Evolution." The student worked '
    'in the Product Department under the supervision of Dmitry Zubretskiy, Chief Product Officer, '
    'in the position of Platform Product Manager responsible for CX automation platforms.',
    space_after=4
)

add_para(
    'During the internship, the student replicated and integrated the \u03c4\u00b2-bench agentic '
    'benchmark as an experimental platform; benchmarked three open-source language models '
    '(Qwen3 30B-A3B, Qwen3.5 Flash, GLM 4.7 Flash) on customer service tasks using OpenRouter; '
    'designed and implemented a Diagnose-Patch-Validate framework for automated AI agent '
    'improvement through prompt and tool evolution; developed a failure taxonomy for agent errors; '
    'conducted systematic experiments across eight conditions (three models, three scales); '
    'and prepared a report and presentation of the findings.',
    space_after=4
)

add_para(
    'Gleb demonstrated a responsible and proactive attitude, working independently with minimal '
    'supervision and consistently meeting deadlines. He showed strong analytical and programming '
    'skills, competence in modern AI infrastructure (LLM APIs, function calling, JSON schemas), '
    'and the ability to present complex technical results clearly. The research produced '
    'quantitative results directly applicable to TargetAI\u2019s product development: a '
    'lightweight methodology for continuous agent improvement without model retraining.',
    space_after=4
)

add_para(
    'The internship program was completed in full. All objectives specified in the individual '
    'PTE assignment were achieved. Based on the results of the internship, I consider '
    'Gleb Dementev professionally suitable for roles in AI product management and applied '
    'AI research. His internship work merits a grade of "Excellent."',
    space_after=12
)

# Signature block
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(
    'Supervisor of Internship from the Organization:\n'
    'Chief Product Officer, TargetAI LLC\n\n'
    '_____________________  /  Zubretskiy D. Y.  /\n\n'
    '"_____"  ___________________  2026\t\t\tM.P.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

output_path = r'C:\Users\Gleb\Work\TargetAI\dev\tau-evo\internship_review_dementev_v2.docx'
doc.save(output_path)
print(f'Review saved to {output_path}')

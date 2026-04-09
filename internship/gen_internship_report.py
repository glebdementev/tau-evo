"""Generate internship report as .docx"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Ensure Times New Roman on all built-in styles
for style_name in ['List Bullet', 'List Number', 'No Spacing', 'Body Text']:
    try:
        s = doc.styles[style_name]
        s.font.name = 'Times New Roman'
        s.font.size = Pt(12)
    except KeyError:
        pass

# Heading styles
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.styles['Heading 1'].font.size = Pt(14)
doc.styles['Heading 2'].font.size = Pt(13)
doc.styles['Heading 3'].font.size = Pt(12)


def add_para(text, bold=False, align=None, size=None, space_after=None, first_line_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    else:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    doc.add_paragraph()  # spacing after table
    return table


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NATIONAL RESEARCH UNIVERSITY\nHIGHER SCHOOL OF ECONOMICS')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Graduate School of Business')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Master\'s Programme "Business Analytics and Big Data Systems"')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INTERNSHIP REPORT')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Optimizing Customer Service Automation with Human-in-the-Loop:\nA Framework for Supervised AI Agent Evolution')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Student: Dementev Gleb Dmitrievich\nGroup: BABDS-242\n\nSupervisor at HSE: Associate Professor Jin S.\nSupervisor at Organization: Zubretskiy D. Y.\n(Chief Product Officer, TargetAI LLC)')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Moscow, 2026')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (placeholder)
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TABLE OF CONTENTS')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

add_para('[To be generated in Word: References > Table of Contents]', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
doc.add_heading('ABSTRACT', level=1)

add_para(
    'This report documents the industrial internship completed at TargetAI LLC (Moscow) '
    'during the period of 12 January to 13 March 2026. The internship was conducted within '
    'the Product Department, where the author holds the position of Platform Product Manager '
    'responsible for the company\u2019s customer experience (CX) automation platforms (TOS1 and TOS2). '
    'The internship supervisor from the organization was Dmitry Zubretskiy, Chief Product Officer; '
    'the academic supervisor was Associate Professor Jin S. of the Graduate School of Business.'
)

add_para(
    'The internship focused on the development and empirical evaluation of a Diagnose-Patch-Validate '
    'framework for automated improvement of AI agent performance on structured customer service '
    'benchmarks. The framework employs a stronger teacher model (Kimi K2.5) to analyze failed '
    'conversation traces of a weaker student model, diagnose root causes of failure, and propose '
    'structured patches to the student\u2019s system prompt, tool schemas, and tool preprocessing logic. '
    'All improvements operate entirely in the input space\u2014no model weights are modified. The design '
    'is motivated by a practical constraint facing enterprises that consume LLMs through APIs: they cannot '
    'modify model weights, yet they need agents that improve continuously from operational experience.'
)

add_para(
    'Experiments were conducted on the airline domain of the \u03c4\u00b2-bench benchmark across three '
    'student models (Qwen3 30B-A3B, Qwen3.5 Flash, GLM 4.7 Flash) and three task-pool scales '
    '(5, 10, and 20 tasks), totalling eight experimental conditions. The framework produced measurable '
    'improvements: Qwen3 30B-A3B\u2019s trial-level pass rate rose from 53% to 73% at 5 tasks and '
    'from 27% to 50% at 10 tasks. Instruction-level patches accounted for 70\u201392% of all successful '
    'fixes across experiments, providing empirical support for the Superficial Alignment Hypothesis in '
    'agentic settings. The results demonstrate that prompt-level knowledge distillation is a viable '
    'alternative to weight-level fine-tuning for enterprise AI agent deployment, while also revealing '
    'its limits: a hard core of 20\u201355% of tasks resists all prompt-level intervention.'
)

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
doc.add_heading('INTRODUCTION', level=1)

add_para(
    'The deployment of large language model (LLM) based agents in enterprise customer service '
    'represents one of the most commercially significant applications of generative AI. Organizations '
    'across industries\u2014financial services, telecommunications, retail, transportation\u2014are '
    'investing in AI-powered agents capable of holding multi-turn conversations, enforcing domain-specific '
    'policies, and invoking back-end tools on behalf of customers. The potential for cost reduction and '
    'service improvement is substantial: a capable AI agent can resolve routine service tickets at a '
    'fraction of the cost and latency of a human operator.'
)

add_para(
    'However, a critical gap separates benchmark performance from enterprise-grade reliability. '
    'Even frontier models such as Claude Opus 4.5 reach approximately 98% on structured tool-calling '
    'benchmarks like \u03c4\u00b2-bench, while production customer-service systems typically require '
    'four-nines reliability (99.99% task success) before organizations are willing to remove the human '
    'safety net (Rabanser et al., 2025). At enterprise scale, even a 2% failure rate translates to tens '
    'of thousands of mishandled customer interactions per month. The problem is compounded by the fact '
    'that deployed agents are static: once their prompt and tool configuration are set, they do not learn '
    'from operational failures. A static agent that fails on a policy edge case on Monday will fail on '
    'the same edge case on Tuesday, and on every subsequent encounter, unless a human engineer manually '
    'diagnoses the failure, rewrites the prompt or tool schema, and re-deploys the agent.'
)

add_para(
    'Three bodies of literature bear on this problem without fully solving it. First, prompt engineering '
    'and tool-use research have improved static agent capability but offer no mechanism for post-deployment '
    'learning (Wei et al., 2022; Yao et al., 2023). Second, alignment techniques such as RLHF and DPO '
    'can shape model behavior but are too heavyweight for rapid, enterprise-specific iteration\u2014they '
    'require training infrastructure, preference data, and access to model weights (Ouyang et al., 2022; '
    'Rafailov et al., 2023). Third, automated prompt optimization has shown that LLMs can improve their '
    'own instructions, but has not been tested on multi-turn tool-calling benchmarks (Agrawal et al., 2025; '
    'Khattab et al., 2023). The gap at the intersection of these three streams is where this work sits.'
)

add_para(
    'This internship was conducted at TargetAI LLC, a Moscow-based B2B company specializing in '
    'customer experience (CX) automation through voice AI agents. The company serves clients across '
    'fintech, retail, telecommunications, insurance, energy, and transportation\u2014industries where '
    'high-volume customer interactions create demand for reliable automation. The work addressed the '
    'challenge of continuous agent improvement by developing and evaluating a framework for automated, '
    'teacher-model-driven prompt and tool evolution. The internship objectives, as defined in the '
    'individual PTE assignment, were as follows:'
)

objectives = [
    'Replicate \u03c4\u00b2-bench agentic tool-calling benchmark results from the original paper to confirm implementation correctness.',
    'Benchmark open-source models of choice (Qwen3 30B-A3B, Qwen3.5 Flash, GLM 4.7 Flash) using organizational computational resources via OpenRouter; report results.',
    'Design a framework that converts teacher-model diagnoses of student failures into iterative prompt, tool schema, and knowledge base improvements.',
    'Evaluate the framework on the \u03c4\u00b2-bench benchmark across multiple scales (5, 10, and 20 tasks); report results.',
    'Analyze the distribution of successful and unsuccessful patches to identify the boundaries of prompt-level intervention.',
    'Prepare a presentation demonstrating methods and findings.',
]

for i, obj in enumerate(objectives, 1):
    add_bullet(f'{i}. {obj}')

add_para(
    'The report is structured as follows. Chapter 1 provides a general overview of TargetAI LLC, '
    'including its history, organizational structure, and product portfolio. Chapter 2 describes '
    'the Product Department and the author\u2019s role within it. Chapter 3 details the internship '
    'activities corresponding to each objective, including benchmark replication, model selection, '
    'framework design, failure taxonomy development, and experimental setup. Chapter 4 presents '
    'the experimental results and key findings across eight experimental conditions. Chapter 5 '
    'discusses the connection between the internship and the master\u2019s thesis. The report '
    'concludes with a summary of outcomes and recommendations for TargetAI and similar organizations.'
)

doc.add_page_break()

# ============================================================
# CHAPTER 1: GENERAL OVERVIEW OF TARGETAI LLC
# ============================================================
doc.add_heading('1. GENERAL OVERVIEW OF TARGETAI LLC', level=1)

doc.add_heading('1.1. Company Profile and History', level=2)

add_para(
    'TargetAI LLC is a Russian technology company founded in 2018 and headquartered in Moscow. '
    'The company specializes in building AI-powered platforms for customer experience automation, '
    'with a focus on voice AI agents for enterprise clients. As of 2026, the company employs '
    'approximately 75 people and generates annual revenue of approximately $3 million USD.'
)

add_para(
    'Since its founding, TargetAI has established itself in the B2B market for conversational AI '
    'solutions, serving clients across multiple industries including financial technology (fintech), '
    'retail and e-commerce, telecommunications, insurance, energy and utilities, and transportation. '
    'The company\u2019s core value proposition is enabling enterprises to automate customer service '
    'interactions through intelligent voice agents that can handle routine inquiries, process requests, '
    'and escalate complex cases to human operators. TargetAI\u2019s technology addresses a growing '
    'market need: as customer service volumes increase and customer expectations for immediacy rise, '
    'enterprises require automated solutions that can operate reliably at scale while maintaining '
    'the quality of human-like interaction.'
)

add_para(
    'The company\u2019s competitive position is built on several pillars: proprietary speech synthesis '
    'technology (TargetSpeak), a modular platform architecture that supports rapid deployment across '
    'industries, and deep domain expertise in voice AI for enterprise use cases. TargetAI has received '
    'coverage in major Russian technology media, including TAdviser and CNews, reflecting its growing '
    'presence in the enterprise AI market.'
)

doc.add_heading('1.2. Organizational Structure', level=2)

add_para(
    'TargetAI operates with a hierarchical organizational structure organized into functional departments. '
    'The key departments include:'
)

add_bullet('Executive Management (CEO)')
add_bullet('Product Department (headed by the Chief Product Officer)')
add_bullet('Engineering and Development')
add_bullet('Sales and Business Development')
add_bullet('Analytics')
add_bullet('Customer Success and Support')

add_para(
    'The company follows a top-down management approach with department heads reporting directly to the '
    'CEO. Strategic and product decisions flow from the CEO through department heads to individual '
    'contributors. The Product Department, where the internship was conducted, is responsible for '
    'defining product strategy, managing platform development, and ensuring alignment between technical '
    'capabilities and market needs. The Engineering and Development department implements product '
    'requirements and maintains the technical infrastructure, while the Analytics department supports '
    'data-driven decision-making across the organization.'
)

doc.add_heading('1.3. Products and Services', level=2)

add_para(
    'TargetAI\u2019s product portfolio centers on AI-powered customer experience automation. '
    'The company\u2019s key offerings include:'
)

add_bullet(
    'TargetAI Platform \u2014 the core platform for creating and deploying voice AI agents capable '
    'of conducting natural-language conversations with customers, handling service requests, and '
    'integrating with enterprise back-end systems. The platform supports multi-turn dialogue, '
    'policy enforcement, and tool invocation\u2014the same capabilities evaluated in this '
    'internship\u2019s benchmark experiments.'
)
add_bullet(
    'TargetSpeak \u2014 a proprietary speech synthesis technology enabling natural-sounding voice '
    'output for AI agents. Low-latency speech synthesis is critical for voice-based customer service, '
    'where response times exceeding 800 milliseconds can raise customer abandonment rates by '
    'approximately 40%.'
)
add_bullet(
    'TargetSkill \u2014 an additional product offering expanding the platform\u2019s capabilities '
    'for specialized use cases and domain-specific agent configurations.'
)
add_bullet(
    'Enterprise Analytics \u2014 analytics tools for monitoring agent performance, conversation '
    'quality, and customer satisfaction metrics. These tools enable the kind of failure detection '
    'and diagnosis that the internship\u2019s evolution framework automates.'
)

add_para(
    'The company positions its solutions as enterprise-ready, offering integration capabilities, '
    'scalability, and compliance features required by large organizations. TargetAI\u2019s client base '
    'spans industries where high-volume customer interactions create demand for automation: banking, '
    'telecom, retail, insurance, and utilities. The diversity of the client base mirrors the '
    'multi-domain structure of the \u03c4\u00b2-bench benchmark used in this internship, which '
    'covers airline, retail, and telecom customer service scenarios.'
)

doc.add_page_break()

# ============================================================
# CHAPTER 2: THE PRODUCT DEPARTMENT AND THE AUTHOR'S ROLE
# ============================================================
doc.add_heading('2. THE PRODUCT DEPARTMENT AND THE AUTHOR\u2019S ROLE', level=1)

doc.add_heading('2.1. Department Functions and Responsibilities', level=2)

add_para(
    'The Product Department at TargetAI is responsible for defining the product vision, managing '
    'the development roadmap, and ensuring that the company\u2019s platforms meet the evolving needs '
    'of enterprise clients. The department operates under the leadership of the Chief Product Officer '
    '(CPO), Dmitry Zubretskiy, who served as the internship supervisor from the organization\u2019s side.'
)

add_para(
    'Key functions of the Product Department include:'
)

add_bullet('Product strategy and roadmap definition for the CX automation platform')
add_bullet('Platform requirements specification and feature prioritization based on client needs')
add_bullet('Coordination with engineering teams for feature development and technical implementation')
add_bullet('Market analysis and competitive intelligence in the conversational AI space')
add_bullet('Client feedback integration and iterative product improvement')
add_bullet('Research and prototyping of new capabilities, including AI agent improvement methodologies')

add_para(
    'The department plays a central role in translating market demands and client feedback into '
    'product capabilities. Given TargetAI\u2019s focus on enterprise CX automation, the Product '
    'Department must balance technical innovation with the reliability and compliance requirements '
    'that enterprise clients demand. This tension\u2014between pushing capability boundaries and '
    'ensuring production-grade reliability\u2014is precisely the challenge that the internship '
    'research addressed.'
)

doc.add_heading('2.2. The Author\u2019s Role', level=2)

add_para(
    'The author holds the position of Platform Product Manager within the Product Department, '
    'responsible for two customer experience platforms: TOS1 and TOS2. These platforms constitute '
    'the core of TargetAI\u2019s CX automation offering, supporting the deployment and management '
    'of voice AI agents across client organizations. In this role, the author operates as an '
    'individual contributor, building product features and conducting research independently '
    'without supervising engineering staff. Product direction and strategic input are received '
    'directly from the CEO.'
)

add_para(
    'The internship activities were directly aligned with the author\u2019s professional responsibilities: '
    'improving the quality and reliability of AI agents deployed on the company\u2019s CX platforms. '
    'The research conducted during the internship\u2014developing a framework for automated agent '
    'improvement through prompt and tool evolution\u2014addresses a core challenge in TargetAI\u2019s '
    'product domain: how to make AI agents learn from failures without costly model retraining. '
    'The framework\u2019s design reflects practical constraints encountered in the author\u2019s '
    'day-to-day work: TargetAI\u2019s platforms consume LLMs through APIs without access to model '
    'weights, making fine-tuning impractical and input-space optimization the only viable path '
    'for rapid improvement.'
)

add_para(
    'The author\u2019s dual role\u2014as both a platform product manager and an internship researcher\u2014'
    'ensured that the framework\u2019s design was grounded in real operational constraints rather than '
    'purely academic considerations. The choice of benchmark (customer service), the emphasis on '
    'API-compatible methods, and the focus on human-readable, auditable patches all reflect requirements '
    'that arise naturally in enterprise CX automation.'
)

doc.add_page_break()

# ============================================================
# CHAPTER 3: INTERNSHIP ACTIVITIES
# ============================================================
doc.add_heading('3. INTERNSHIP ACTIVITIES', level=1)

doc.add_heading('3.1. Goals and Objectives', level=2)

add_para(
    'The internship was structured around the development and empirical evaluation of a framework '
    'for automated AI agent improvement. The central research question guiding the work was: '
    'How can AI agent performance on structured benchmarks be improved through automated, '
    'teacher-model-driven prompt and tool evolution?'
)

add_para(
    'Three sub-questions guided the investigation:'
)

add_bullet(
    'Which failure modes in multi-turn tool-calling tasks are most responsive to prompt and '
    'tool-schema edits, and which resist them?'
)
add_bullet(
    'What is the scaling behavior of the evolution framework: does fix rate hold as the task pool '
    'grows, or does it degrade?'
)
add_bullet(
    'Can prompt-level patches generated by a teacher model close a meaningful fraction of the '
    'performance gap between a weak student and a frontier ceiling, without any weight updates?'
)

add_para(
    'These questions were motivated by a practical need at TargetAI: the company deploys AI agents '
    'that fail on edge cases, and the current remediation process requires manual diagnosis and prompt '
    'rewriting by engineers. Automating this loop\u2014from failure detection through diagnosis to '
    'validated fix\u2014would reduce operational costs and accelerate the improvement cycle.'
)

doc.add_heading('3.2. Benchmark Replication and Baseline Evaluation', level=2)

add_para(
    'The first phase of the internship involved replicating the \u03c4\u00b2-bench agentic '
    'tool-calling benchmark (Barres et al., 2025). \u03c4\u00b2-bench is an extension of \u03c4-bench '
    '(Yao et al., 2024) that captures the full tool-agent-user interaction loop as a binary pass/fail '
    'criterion. In \u03c4\u00b2-bench, a simulated user reveals information incrementally across turns '
    'while the agent must follow domain-specific policies, call the correct tools with correct parameters, '
    'and modify database state accordingly. A task passes only when every sub-criterion in the reward '
    'rubric scores 1.0\u2014the strict binary pass\u00b9 metric standard in \u03c4-bench publications.'
)

add_para(
    'The benchmark was selected over alternatives (AgentBench, SWE-bench, GAIA, ToolBench) because '
    'it is the only benchmark that combines all four elements required for evaluating customer service '
    'agents: multi-turn conversations with an LLM-simulated user providing partial information, '
    'domain-specific policies, tool-calling APIs that modify database state, and a reliability metric. '
    'The benchmark spans three customer-service domains:'
)

add_table(
    ['Domain', 'Tasks', 'Description'],
    [
        ['Airline', '50', 'Flight reservations, cancellations, upgrades, baggage policies'],
        ['Retail', '114', 'Order management, returns, exchanges, account issues'],
        ['Telecom', '114', 'Mobile plans, data usage, billing, service changes'],
    ]
)

add_para(
    '\u03c4\u00b2-bench was integrated as a git submodule pinned to commit 37bfc31 (based on '
    'tag v0.1.1), installed as an editable Python package. No modifications were made to the upstream '
    'codebase; all integration occurred through the benchmark\u2019s public API: the RunConfig data '
    'model, the run_domain() function, the agent registry, and the Tool and Environment classes. '
    'Because nothing in the upstream code was changed, benchmark results are directly comparable '
    'to published baselines.'
)

doc.add_heading('3.3. Model Selection and Benchmarking', level=2)

add_para(
    'Three open-source student models were selected for benchmarking, all accessed through OpenRouter '
    'using a single API key. The selection reflects a deliberate trade-off: the primary student model '
    'needed non-trivial \u03c4\u00b2-bench scores but had to be weak enough relative to the frontier '
    'that meaningful headroom for improvement existed, and cheap enough for the many evaluation runs '
    'that the iterative evolution process requires.'
)

add_table(
    ['Role', 'Model', 'Description', 'Access Method'],
    [
        ['Student (primary)', 'Qwen3 30B-A3B', 'MoE, 30B total / 3.3B active params', 'litellm via \u03c4\u00b2-bench'],
        ['Student (strong)', 'Qwen3.5 Flash', 'Newer-gen, stronger instruction following', 'litellm via \u03c4\u00b2-bench'],
        ['Student (cross-family)', 'GLM 4.7 Flash', 'Dense model, GLM family', 'litellm via \u03c4\u00b2-bench'],
        ['User simulator', 'Qwen3 30B-A3B', 'Follows scripted scenarios', 'litellm via \u03c4\u00b2-bench'],
        ['Teacher', 'Kimi K2.5', '~1T total / 32B active, 256K context', 'OpenAI client direct'],
    ]
)

add_para(
    'Qwen3 30B-A3B (Qwen, 2025) is a Mixture-of-Experts Transformer with 30.5 billion total parameters '
    'but only 3.3 billion active per token. It employs 128 experts with top-8 routing across 48 layers. '
    'Despite activating barely 10% of its parameters per forward pass, it leads the Berkeley Function '
    'Calling Leaderboard (BFCL v3). The MoE architecture makes it fast and inexpensive via API while '
    'handling tool-calling and multi-turn dialogue competently. Qwen3.5 Flash is a newer-generation model '
    'with substantially stronger instruction-following and tool-use capabilities, included to test how '
    'baseline student capability affects the framework\u2019s value proposition. GLM 4.7 Flash is a dense '
    'model from the GLM family with distinct architectural characteristics, included to test whether the '
    'framework generalizes beyond a single model family.'
)

add_para(
    'The teacher model, Kimi K2.5 (Moonshot AI, 2026), is a visual-agentic extension of Kimi K2 with '
    'approximately one trillion total parameters and 32 billion active per token. Its 256K-token context '
    'window accommodates full conversation traces (often 5\u201315 turns with tool calls and results), '
    'the system prompt, all tool schemas, and task requirements in a single prompt. The teacher was chosen '
    'for its significantly stronger performance than the student on target domains, strong tool-calling '
    'comprehension, a long context window, architectural independence from the student (Moonshot AI, not '
    'Alibaba), and cost-effectiveness for hundreds of reflection calls via OpenRouter.'
)

doc.add_heading('3.4. Design of the Evolution Framework', level=2)

add_para(
    'The core contribution of the internship was the design and implementation of a Diagnose-Patch-Validate '
    'framework for automated agent improvement. The framework implements a teacher-driven variant of '
    'reflective prompt evolution: a stronger teacher model analyzes a weaker student model\u2019s failed '
    'conversation traces, diagnoses root causes, and proposes structured patches. The closest methodological '
    'precedent is the GEPA framework (Agrawal et al., 2025), which uses natural-language reflection to '
    'diagnose failures and propose prompt mutations. The present framework departs from GEPA in three '
    'respects: it operates on three distinct patch surfaces (not just prompt text), every patch is '
    'validated by re-running the student before merging, and it targets a multi-turn tool-calling '
    'benchmark rather than single-turn reasoning tasks.'
)

add_para(
    'The three patch surfaces are:'
)

add_bullet(
    'Prompt patches \u2014 additions or clarifications of behavioral rules in the system prompt. '
    'These are the lightest-weight intervention: a rule such as \u201calways verify the customer\u2019s '
    'identity before processing a cancellation\u201d is simply appended to the system prompt. '
    'The Superficial Alignment Hypothesis (Zhou et al., 2023) suggests this should work: alignment '
    'primarily teaches style and format, which prompt text can supply.'
)
add_bullet(
    'Tool schema patches \u2014 modifications to the JSON schemas that define how the agent calls '
    'each tool. Common modifications include clarifying parameter descriptions (e.g., adding '
    '\u201cmust start with #\u201d to a reservation_id field), expanding tool descriptions, and '
    'adding constraint notes. After each edit, the JSON string is parsed to ensure syntactic validity.'
)
add_bullet(
    'Tool preprocessors \u2014 sandboxed Python functions that transform tool-call arguments before '
    'execution. Some formatting errors persist even when the prompt and schema are clear: the model '
    'understands the requirement but occasionally gets it wrong due to tokenization or sampling '
    'artifacts. A preprocessor catches such errors at the tool-call boundary. Preprocessors are '
    'sandboxed: a static analysis pass rejects forbidden constructs (imports, eval, exec, file I/O).'
)

add_para(
    'The framework operates through two nested loops. The Outer Loop iterates over the full task set '
    'in successive sweeps. In each sweep, every task is evaluated with the current evolved state; tasks '
    'that fail are passed to the inner loop for repair. Re-evaluating all tasks every sweep\u2014rather '
    'than dropping attempted tasks\u2014is deliberate: merged patches from multiple independent teacher '
    'sessions can interact, and re-evaluation catches regressions introduced by the merge step.'
)

add_para(
    'The Inner Loop processes each failed task through a teacher session. The teacher model receives '
    'a comprehensive prompt containing the agent\u2019s current system prompt, all tool schemas, '
    'the full failed conversation trace, the task requirements, and the reward breakdown. It diagnoses '
    'the root cause, classifies the failure, and calls patch tools to propose modifications. The patch '
    'is validated through re-simulation: the student is re-run on the same task with the proposed '
    'patch applied. A fix is accepted only if the task passes unanimously\u2014all trials achieve '
    'a perfect reward of 1.0. If validation fails, all patches are reverted and the teacher retries '
    'with a different approach.'
)

add_para(
    'The inner loop employs a two-phase escalation strategy. In Phase 1 (teaching), the teacher can '
    'only modify the prompt and tool schemas. If Phase 1 exhausts its attempts without fixing the task, '
    'Phase 2 (guardrails) unlocks the tool preprocessor editing capability, allowing the teacher to add '
    'defensive input coercion. This staged approach ensures that lighter-weight interventions are '
    'attempted first, reflecting the empirical finding that 70\u201392% of fixes require only '
    'instruction-level patches.'
)

add_para(
    'The framework was implemented in Python using the \u03c4\u00b2-bench public API, the OpenAI '
    'Python SDK for teacher model calls, and litellm for model routing. All parameters are fixed '
    '(seed = 42), task IDs are locked after the first evaluation, and the complete evolution state '
    '(patches, traces, evaluation results) is serialized to JSON after each sweep for full '
    'reproducibility. The source code is publicly available.'
)

doc.add_heading('3.5. Failure Taxonomy Development', level=2)

add_para(
    'As part of the framework development, a failure taxonomy was created to classify the types '
    'of errors observed in student model conversation traces. The teacher classifies each failure '
    'into one of four categories as part of its diagnostic output:'
)

add_table(
    ['Category', 'Description', 'Example'],
    [
        ['TOOL_MISUSE', 'Wrong tool, wrong parameters, missing tool call', 'Using get_flight_details instead of get_reservation_details'],
        ['POLICY_VIOLATION', 'Skipped validation step or broke a constraint', 'Cancelling without checking refund eligibility'],
        ['REASONING_ERROR', 'Incorrect assumption, incomplete plan', 'Assuming a flight is direct when it has connections'],
        ['COMMUNICATION_ERROR', 'Confusing message, failed to guide user', 'Not explaining applicable fees to the customer'],
    ]
)

add_para(
    'The taxonomy distinguishes between failure modes that are responsive to prompt-level intervention '
    'and those that resist it. Policy violations and communication errors are typically addressable '
    'through instruction patches\u2014adding explicit rules to the system prompt. Tool misuse may '
    'require schema-level patches to clarify parameter requirements. Reasoning errors constitute the '
    '\u201chard core\u201d of resistant tasks: they require capabilities that the student model does '
    'not possess and that cannot be injected through prompt text. This classification informed the '
    'framework\u2019s three-tier escalation strategy and provided a basis for the per-category analysis '
    'of which failure types respond to evolution.'
)

doc.add_heading('3.6. Experimental Setup', level=2)

add_para(
    'The framework was evaluated across increasing task-pool sizes to characterize scaling behavior. '
    'The experimental design follows a design-science approach (Hevner et al., 2004): the primary '
    'contribution is a software artifact\u2014the prompt evolution framework\u2014and the primary '
    'evaluation is an empirical measurement of its effect on a standardized benchmark. Three '
    'experimental conditions were defined, forming a floor\u2013intervention\u2013ceiling comparison:'
)

add_table(
    ['Condition', 'Description', 'Purpose'],
    [
        ['Baseline (B)', 'Student model with unmodified default prompt and tools', 'Performance floor'],
        ['Evolved (K)', 'Student model with evolved configuration from the framework', 'Intervention effect'],
        ['Frontier (F)', 'Teacher model (Kimi K2.5) running directly as agent', 'Performance ceiling'],
    ]
)

add_para(
    'The three conditions allow computing a gap closure metric: what fraction of the baseline-to-frontier '
    'performance gap is closed by prompt-level evolution alone. This normalization provides context '
    'that raw pass rates cannot: a pass rate of 60% means nothing without knowing where the ceiling is.'
)

add_para(
    'Each experiment ran the evolution loop for up to three sweeps with up to two retries per '
    'failed task per sweep. Every task was evaluated with three trials; a task was considered '
    'passing if it passed in at least two of three trials (majority vote). The experiments spanned '
    'three scales to test the scaling hypothesis:'
)

add_table(
    ['Scale', 'Task IDs', 'Models Evaluated'],
    [
        ['5 tasks', '0, 1, 3, 4, 5', 'All three student models'],
        ['10 tasks', '0, 1, 3, 4, 5, 7, 9, 10, 11, 12', 'All three student models'],
        ['20 tasks', '0, 1, 3, 4, 5, 7, 9, 10\u201334 (selected)', 'Qwen3 30B-A3B, Qwen3.5 Flash'],
    ]
)

add_para(
    'The scaling sequence is deliberate. If the evolution framework captures task-specific fixes that '
    'do not generalize, then gains should be largest when the task set is smallest\u2014each fix '
    'represents a larger share of the total\u2014and should diminish as the denominator grows. '
    'The multi-model comparison tests a complementary question: if a stronger student already passes '
    'most tasks at baseline, does the framework still provide value? GLM 4.7 Flash was dropped from '
    'the 20-task experiment due to poor performance at 10 tasks, where the framework produced zero '
    'net improvement and actively degraded baseline performance.'
)

doc.add_page_break()

# ============================================================
# CHAPTER 4: RESULTS AND KEY FINDINGS
# ============================================================
doc.add_heading('4. RESULTS AND KEY FINDINGS', level=1)

doc.add_heading('4.1. Cross-Scale Results', level=2)

add_para(
    'This section presents the aggregated results across all experimental conditions. Detailed '
    'per-task trajectories, per-sweep heatmaps, and individual fix attempt logs are documented '
    'in the thesis and supporting materials.'
)

add_para(
    'Table 1 summarizes the cross-scale results for the primary student model, Qwen3 30B-A3B. '
    'The baseline trial rate declines as harder tasks enter the pool (53% \u2192 27% \u2192 22%), '
    'while the framework produces consistent but diminishing absolute improvements.'
)

add_table(
    ['Tasks', 'Baseline trial rate', 'Best trial rate', 'Improvement (pp)', 'Fix rate on failing'],
    [
        ['5', '53% (8/15)', '73% (11/15)', '+20', '100% (4/4)'],
        ['10', '27% (8/30)', '50% (15/30)', '+23', '71% (5/7)'],
        ['20', '22% (13/60)', '33% (20/60)', '+11', '53% (8/15)'],
    ]
)

add_para(
    'Table 2 summarizes results for Qwen3.5 Flash, the stronger student model. At 5 tasks, '
    'no evolution is needed\u2014the model achieves a perfect pass rate out of the box. At larger '
    'scales, the framework produces substantial improvements, with the stronger student achieving '
    'higher absolute ceilings and higher fix rates than Qwen3 30B-A3B.'
)

add_table(
    ['Tasks', 'Baseline trial rate', 'Best trial rate', 'Improvement (pp)', 'Fix rate on failing'],
    [
        ['5', '100% (15/15)', '100% (15/15)', '0', 'N/A (no failures)'],
        ['10', '60% (18/30)', '80% (24/30)', '+20', '80% (4/5)'],
        ['20', '47% (28/60)', '58% (35/60)', '+11', '45% (5/11)'],
    ]
)

add_para(
    'Table 3 summarizes results for GLM 4.7 Flash, which presents a cautionary tale. At 5 tasks, '
    'the framework produces a strong peak improvement (+26pp trial, +40pp majority at sweep 2), '
    'but gains are entirely erased by sweep 3 due to catastrophic regression. At 10 tasks, the '
    'framework fails completely: zero genuinely failing tasks are fixed, and baseline performance '
    'degrades monotonically across sweeps.'
)

add_table(
    ['Tasks', 'Baseline trial rate', 'Best trial rate', 'Improvement (pp)', 'Fix rate on failing'],
    [
        ['5', '47% (7/15)', '73% (11/15)', '+26', '67% (2/3)'],
        ['10', '50% (15/30)', '50% (15/30)', '0', '0% (0/4)'],
    ]
)

doc.add_heading('4.2. Cross-Model Comparison', level=2)

add_para(
    'Table 4 presents a cross-model comparison at matched scales, revealing that the set of '
    'unfixable tasks is model-dependent, not task-intrinsic.'
)

add_table(
    ['Metric', 'Qwen3 30B-A3B (10 tasks)', 'Qwen3.5 Flash (10 tasks)', 'GLM 4.7 Flash (10 tasks)'],
    [
        ['Baseline majority', '30%', '50%', '60%'],
        ['Best majority', '50%', '80%', '60%'],
        ['Fix rate (failing)', '71%', '80%', '0%'],
        ['Unfixable tasks', '4 (7, 9, 11, 12)', '1 (7)', '4 (7, 9, 11, 12)'],
    ]
)

add_para(
    'The most important finding from the cross-model comparison is that tasks unfixable for weaker '
    'students can become fixable for stronger ones. Tasks 9, 11, and 12\u2014part of Qwen3 30B-A3B\u2019s '
    '\u201chard core\u201d of unfixable tasks\u2014are all fixed by instruction patches on Qwen3.5 Flash. '
    'This confirms that these tasks are not inherently beyond prompt-level correction; rather, the weaker '
    'model lacked the baseline capability to execute even well-specified instructions. Only Task 7 remains '
    'genuinely resistant across all three models, representing a structural challenge that neither model '
    'can handle through prompt-level intervention alone.'
)

add_para(
    'At 20 tasks, only two models are compared. Qwen3.5 Flash achieves 65% majority-vote pass rate '
    'versus Qwen3 30B-A3B\u2019s 30%, with 5 unfixable tasks versus 11. The framework amplifies '
    'capability differences rather than equalizing them: investing in a stronger base model yields '
    'compounding returns when combined with prompt evolution.'
)

doc.add_heading('4.3. Key Findings', level=2)

add_para(
    'The experimental evaluation produced seven principal findings that characterize the framework\u2019s '
    'effectiveness and limitations.'
)

add_para(
    'Finding 1: Prompt evolution produces measurable improvement. Across both Qwen models, the '
    'framework consistently raised pass rates by 11\u201323 percentage points without any weight '
    'updates, gradient computation, or preference data\u2014only through structured edits to the '
    'student\u2019s system prompt and tool interface. These improvements came from a lightweight, '
    'input-space-only mechanism compatible with API-served models.'
)

add_para(
    'Finding 2: Instruction-tier patches dominate. Across all experiments, 70\u201392% of successful '
    'fixes were prompt-level instruction patches (additions or clarifications of behavioral rules in '
    'the system prompt), while the remaining fixes were guardrail-tier patches involving tool schema '
    'constraints or input preprocessors. This 71/29 split held precisely constant across the first two '
    'experiments despite doubling the task pool size. The dominance of instruction patches provides '
    'empirical support for the Superficial Alignment Hypothesis (Zhou et al., 2023): student models '
    'already possess the capability to perform correctly but lack explicit knowledge that they should '
    'follow a particular policy. Stating the requirement in plain text is sufficient to unlock the '
    'capability in most fixable cases.'
)

add_para(
    'Finding 3: Absolute gain is stable but fix rate declines with scale. The framework produced '
    'roughly the same absolute improvement (~20\u201323 pp) at 5 and 10 tasks, but the fix rate '
    'dropped from 100% to 56\u201380%. At 20 tasks, improvement halved to ~11 pp. The framework\u2019s '
    'capacity ceiling is approximately 5\u20138 fixable tasks per model, regardless of pool size: '
    'additional tasks primarily contribute unfixable failures that consume teacher effort without '
    'yielding successful patches.'
)

add_para(
    'Finding 4: A hard core of resistant tasks emerges. Tasks requiring multi-step reasoning under '
    'uncertainty, implicit policy interpretation, or complex state tracking resisted all fix attempts '
    'across multiple sweeps. At 10 tasks, four tasks (7, 9, 11, 12) consumed over 150 messages, '
    '61 tool calls, and 36 minutes of teacher compute in a single sweep without yielding a single '
    'viable patch. These tasks define the boundary of what prompt-level intervention can achieve.'
)

add_para(
    'Finding 5: The framework saturates rapidly. In all experiments, diminishing returns set in '
    'after the first two sweeps, with sweep 3 producing zero new fixes and occasional regressions. '
    'The practical implication: the framework\u2019s value is concentrated in the first one or two '
    'passes, making it efficient for deployment but limited in cumulative impact.'
)

add_para(
    'Finding 6: Patch interference is real. Accumulated patches can degrade performance on '
    'previously passing tasks\u2014an effect analogous to catastrophic forgetting in continual '
    'learning (Luo et al., 2023), but operating in prompt space rather than weight space. The '
    'severity varies by model: Qwen3 30B-A3B shows mild regression, Qwen3.5 Flash shows severe '
    'discrete regression at 10 tasks (\u221217pp trial rate), and GLM 4.7 Flash shows catastrophic '
    'collapse (\u221240pp majority at 5 tasks). Production deployments would need patch management '
    'discipline: versioning, regression testing, and potentially patch retirement.'
)

add_para(
    'Finding 7: Student model strength reshapes the framework\u2019s utility. The stronger student '
    '(Qwen3.5 Flash) fixes tasks that the weaker model cannot, achieving higher absolute ceilings '
    '(65% vs. 30% at 20 tasks by majority vote). GLM 4.7 Flash demonstrated that the framework can '
    'be actively harmful when applied to a model that cannot reliably execute patched instructions\u2014'
    'a negative result that, to our knowledge, has not been systematically documented in the prompt '
    'optimization literature. The framework has a minimum student capability threshold below which '
    'patches cause net harm.'
)

doc.add_heading('4.4. Fix Tier Breakdown', level=2)

add_para(
    'Table 5 presents the distribution of fix types across all eight experimental conditions, '
    'confirming instruction-level patching as the dominant mechanism of improvement.'
)

add_table(
    ['Experiment', 'Total fixes', 'Instruction', 'Tools', 'Guardrail'],
    [
        ['Qwen3 30B, 5 tasks', '7', '5 (71%)', '0 (0%)', '2 (29%)'],
        ['Qwen3 30B, 10 tasks', '7', '5 (71%)', '0 (0%)', '2 (29%)'],
        ['Qwen3 30B, 20 tasks', '10', '7 (70%)', '2 (20%)', '1 (10%)'],
        ['Qwen3.5 Flash, 10 tasks', '5', '4 (80%)', '0 (0%)', '1 (20%)'],
        ['Qwen3.5 Flash, 20 tasks', '12', '11 (92%)', '0 (0%)', '1 (8%)'],
        ['GLM 4.7, 5 tasks', '4', '3 (75%)', '0 (0%)', '1 (25%)'],
        ['GLM 4.7, 10 tasks', '4', '3 (75%)', '0 (0%)', '1 (25%)'],
    ]
)

add_para(
    'Two trends stand out. First, the stronger student (Qwen3.5 Flash) shows a progressively higher '
    'instruction-tier share as scale increases (80% at 10 tasks, 92% at 20 tasks), suggesting that '
    'stronger instruction-following capability reduces the need for guardrail or tool-level interventions. '
    'Second, tool-schema patches appear only for Qwen3 30B-A3B at 20 tasks (2 of 10 fixes)\u2014neither '
    'Qwen3.5 Flash nor GLM 4.7 Flash required tool-level intervention at any scale. This confirms that '
    'the escalation tiers are correctly ordered by generality: most failures can be addressed at the '
    'lightest level of intervention.'
)

doc.add_page_break()

# ============================================================
# CHAPTER 5: CONNECTION TO THE THESIS
# ============================================================
doc.add_heading('5. CONNECTION BETWEEN THE INTERNSHIP AND THE THESIS', level=1)

add_para(
    'The internship activities were directly aligned with the master\u2019s thesis titled '
    '\u201cOptimizing Customer Service Automation with Human-in-the-Loop: A Framework for '
    'Supervised AI Agent Evolution.\u201d The thesis and internship share the same research question, '
    'methodology, and experimental apparatus. The internship provided both the practical motivation '
    'and the computational infrastructure for the thesis research.'
)

add_para(
    'The internship contributed to the thesis in the following specific ways:'
)

add_bullet(
    'Benchmark infrastructure: The replication and integration of \u03c4\u00b2-bench during '
    'the internship provided the experimental platform used for all thesis experiments. The '
    'integration was non-trivial, requiring adaptation of the benchmark\u2019s API to work with '
    'OpenRouter model routing and the evolution framework\u2019s parallel execution architecture.'
)
add_bullet(
    'Framework development: The Diagnose-Patch-Validate framework designed and implemented during '
    'the internship constitutes the core artifact of the thesis. The framework\u2019s architecture\u2014'
    'the outer loop, inner loop, teacher session, three patch surfaces, and two-phase escalation\u2014'
    'was developed iteratively during the internship period.'
)
add_bullet(
    'Empirical results: All experimental results reported in the thesis (eight conditions across three '
    'models and three scales) were generated during the internship period using TargetAI\u2019s '
    'computational resources via OpenRouter API access.'
)
add_bullet(
    'Practical validation: The internship context at a CX automation company provided real-world '
    'relevance to the research, grounding the framework\u2019s design in the practical constraints '
    'of enterprise AI deployment\u2014API-only model access, latency requirements, auditability needs.'
)

add_para(
    'The thesis extends the internship work in several directions. It includes a comprehensive '
    'literature review (approximately 90 references) positioning the framework within the broader '
    'landscape of agentic benchmarking, prompt optimization, knowledge distillation, and human-in-the-loop '
    'alignment. It provides deeper theoretical analysis, connecting the empirical findings to the '
    'Superficial Alignment Hypothesis (Zhou et al., 2023) and drawing analogies to catastrophic '
    'forgetting in the continual learning literature (Luo et al., 2023). It also frames the framework '
    'as a form of prompt-level knowledge distillation, extending Hinton et al.\u2019s (2015) weight-level '
    'distillation paradigm into the input space. Finally, the thesis articulates the framework\u2019s '
    'contribution as a resolution of a System 1/System 2 tension in enterprise AI: leveraging slow, '
    'deliberate reasoning at design time (the teacher) to improve fast, automatic execution at run time '
    '(the student)\u2014mirroring how human organizations operate when experts design procedures for '
    'frontline staff.'
)

doc.add_page_break()

# ============================================================
# CONCLUSIONS AND RECOMMENDATIONS
# ============================================================
doc.add_heading('CONCLUSIONS AND RECOMMENDATIONS', level=1)

doc.add_heading('Conclusions', level=2)

add_para(
    'The internship achieved its primary objectives. The \u03c4\u00b2-bench benchmark was successfully '
    'replicated and integrated as an experimental platform. Three open-source models (Qwen3 30B-A3B, '
    'Qwen3.5 Flash, GLM 4.7 Flash) were benchmarked on the airline domain via OpenRouter. A '
    'Diagnose-Patch-Validate framework was designed, implemented in Python, and empirically evaluated '
    'across eight experimental conditions spanning three student models and three task-pool scales '
    '(5, 10, and 20 tasks). The distribution of successful and unsuccessful patches was analyzed, '
    'identifying the boundaries of prompt-level intervention. A presentation demonstrating methods '
    'and findings was prepared.'
)

add_para(
    'The principal conclusions are:'
)

add_bullet(
    'Prompt-level evolution is a viable mechanism for improving AI agent performance without '
    'weight modifications. The framework produced consistent improvements of 11\u201323 percentage '
    'points in trial pass rates across Qwen-family models, using only structured edits to the '
    'student\u2019s system prompt and tool interface. This approach is compatible with API-served '
    'models and requires no GPU infrastructure beyond inference.'
)
add_bullet(
    'Instruction-level patches are the dominant lever, accounting for 70\u201392% of all successful '
    'fixes across eight experiments. This finding has direct practical implications: organizations '
    'should invest in thorough policy documentation and structured prompt engineering before '
    'considering fine-tuning. The finding also provides empirical support for the Superficial '
    'Alignment Hypothesis in multi-turn agentic settings.'
)
add_bullet(
    'The framework\u2019s effectiveness is contingent on the student model\u2019s ability to absorb '
    'and execute patches reliably. GLM 4.7 Flash demonstrated that prompt evolution can be actively '
    'harmful when applied to an incompatible model. Model compatibility testing is essential before '
    'deployment.'
)
add_bullet(
    'Prompt-level patching has a natural ceiling. A hard core of 20\u201355% of tasks resists '
    'all prompt-level intervention, requiring complementary approaches\u2014stronger base models, '
    'fine-tuning, architectural improvements, and human escalation\u2014for enterprise-grade '
    'reliability. The framework is best understood as one layer in a broader improvement strategy.'
)

doc.add_heading('Recommendations', level=2)

add_para(
    'Based on the internship findings, the following recommendations are offered for TargetAI '
    'and similar organizations deploying AI agents in customer service:'
)

add_bullet(
    'Adopt a prompt evolution pipeline as a first-line improvement mechanism. The framework\u2019s '
    'lightweight, input-space-only approach is compatible with API-served models and existing '
    'change management workflows. Patches are human-readable text edits that can be reviewed, '
    'versioned, approved, and rolled back through standard processes.'
)
add_bullet(
    'Invest in the strongest feasible base model before applying prompt evolution. The framework '
    'amplifies capability differences: a stronger student model achieves both higher absolute '
    'ceilings and higher fix rates. Qwen3.5 Flash reached 65% majority at 20 tasks versus '
    'Qwen3 30B-A3B\u2019s 30%\u2014a difference driven by the student\u2019s ability to execute '
    'patched instructions, not by the quality of the patches themselves.'
)
add_bullet(
    'Implement patch management discipline. Version patches, run regression tests against held-out '
    'task sets, and limit evolution to two sweeps to avoid patch interference and regression. '
    'The observed regression patterns differ by model and require model-specific strategies.'
)
add_bullet(
    'Conduct model compatibility testing before deploying the framework with a new student model. '
    'A brief pilot evaluation at small scale (5 tasks) can reveal whether a model can absorb '
    'patches without destabilization, avoiding the GLM 4.7 Flash failure pattern.'
)
add_bullet(
    'Explore hybrid prompt-and-weight approaches for resistant tasks. Combine prompt evolution '
    'for the accessible long tail of policy-encodable failures with lightweight fine-tuning (LoRA) '
    'for the hard core of structurally resistant failures. The two approaches may be complementary, '
    'addressing different failure modes.'
)

doc.add_page_break()

# ============================================================
# BIBLIOGRAPHY
# ============================================================
doc.add_heading('BIBLIOGRAPHY', level=1)

refs = [
    'Agrawal, M., Saha, S., & Chen, Z. (2025). Generalized Evolutionary Prompt Augmentation for Unified Optimization of Language Model Prompts. In Proceedings of the International Conference on Learning Representations (ICLR 2026).',
    'Bai, Y., Kadavath, S., Kundu, S., et al. (2022). Constitutional AI: Harmlessness from AI Feedback. arXiv preprint arXiv:2212.08073.',
    'Barres, L., Yao, S., & Chen, D. (2025). \u03c4\u00b2-bench: Benchmarking Tool-Agent-User Interaction in Real-World Domains. arXiv preprint.',
    'Fernando, C., Banarse, D., Michalewski, H., et al. (2023). PromptBreeder: Self-Referential Self-Improvement via Prompt Evolution. arXiv preprint arXiv:2309.16797.',
    'Guo, Q., Wang, R., Guo, J., et al. (2023). Connecting Large Language Models with Evolutionary Algorithms Yields Powerful Prompt Optimizers. arXiv preprint arXiv:2309.08532.',
    'Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design Science in Information Systems Research. MIS Quarterly, 28(1), 75\u2013105.',
    'Hinton, G., Vinyals, O., & Dean, J. (2015). Distilling the Knowledge in a Neural Network. arXiv preprint arXiv:1503.02531.',
    'Hu, E. J., Shen, Y., Wallis, P., et al. (2022). LoRA: Low-Rank Adaptation of Large Language Models. In Proceedings of the International Conference on Learning Representations (ICLR 2022).',
    'Huang, K., et al. (2025). CRMArena: Understanding the Capacity of LLM Agents to Perform Professional CRM Tasks in Simulated Environments. arXiv preprint.',
    'Kapoor, S., et al. (2024). AI Agents That Matter. arXiv preprint.',
    'Khattab, O., Santhanam, K., Li, X. L., et al. (2023). DSPy: Compiling Declarative Language Model Calls into Self-Improving Pipelines. arXiv preprint arXiv:2310.03714.',
    'Kimi Team. (2025). Kimi K2: A Frontier Mixture-of-Experts Model. Moonshot AI Technical Report.',
    'Kimi Team. (2026). Kimi K2.5: Visual-Agentic Extension of K2. Moonshot AI Technical Report.',
    'Liu, X., et al. (2023). AgentBench: Evaluating LLMs as Agents. In Proceedings of the International Conference on Learning Representations (ICLR 2024).',
    'Luo, Y., et al. (2023). An Empirical Study of Catastrophic Forgetting in Large Language Models during Continual Fine-Tuning. arXiv preprint.',
    'Ouyang, L., Wu, J., Jiang, X., et al. (2022). Training Language Models to Follow Instructions with Human Feedback. In Advances in Neural Information Processing Systems (NeurIPS 2022).',
    'Qin, Y., et al. (2023). ToolLLM: Facilitating Large Language Models to Master 16000+ Real-World APIs. arXiv preprint arXiv:2307.16789.',
    'Qwen Team. (2025). Qwen3 Technical Report. Alibaba Cloud.',
    'Rabanser, S., et al. (2025). On Agent Reliability. arXiv preprint.',
    'Rafailov, R., Sharma, A., Mitchell, E., et al. (2023). Direct Preference Optimization: Your Language Model Is Secretly a Reward Model. In Advances in Neural Information Processing Systems (NeurIPS 2023).',
    'Sclar, M., et al. (2023). Quantifying Language Models\u2019 Sensitivity to Spurious Features in Prompt Design. arXiv preprint.',
    'Shinn, N., et al. (2023). Reflexion: Language Agents with Verbal Reinforcement Learning. In Advances in Neural Information Processing Systems (NeurIPS 2023).',
    'Wei, J., et al. (2022). Chain-of-Thought Prompting Elicits Reasoning in Large Language Models. In Advances in Neural Information Processing Systems (NeurIPS 2022).',
    'Yao, S., Zhao, J., Yu, D., et al. (2023). ReAct: Synergizing Reasoning and Acting in Language Models. In Proceedings of the International Conference on Learning Representations (ICLR 2023).',
    'Yao, S., et al. (2024). \u03c4-bench: A Benchmark for Tool-Agent-User Interaction in Multi-Turn Conversations. In Advances in Neural Information Processing Systems (NeurIPS 2024).',
    'Yuksekgonul, M., et al. (2024). TextGrad: Automatic \u201cDifferentiation\u201d via Text. arXiv preprint.',
    'Zhou, C., Liu, P., Xu, P., et al. (2023). LIMA: Less Is More for Alignment. In Advances in Neural Information Processing Systems (NeurIPS 2023).',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.hanging_indent = Cm(1.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'{i}. {ref}')
    run.font.size = Pt(12)

# ============================================================
# SAVE
# ============================================================
output_path = r'C:\Users\Gleb\Work\TargetAI\dev\tau-evo\internship_report_dementev.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
